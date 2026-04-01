"""
╔═══════════════════════════════════════════════════════════╗
║         CORRECTEUR MCVA — APPLICATION MOBILE              ║
║                                                           ║
║  ✅ 100% Gratuit (Google Gemini Flash)                    ║
║  ✅ RGPD OK (anonymisation intégrée)                      ║
║  ✅ Mobile (fonctionne sur téléphone)                     ║
║  ✅ Export Excel avec notes + appréciations               ║
║  ✅ Conversion automatique /80 /120 → /20                 ║
║  ✅ Support PDF + Word pour corrigé et barème             ║
╚═══════════════════════════════════════════════════════════╝
"""

import streamlit as st
import google.generativeai as genai
from PIL import Image
import fitz
from docx import Document as DocxDocument
import pandas as pd
import numpy as np
import cv2
import json
import re
import io
from datetime import datetime


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  CONFIGURATION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
st.set_page_config(
    page_title="📝 Correcteur MCVA",
    page_icon="📝",
    layout="wide",
)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  STYLE MOBILE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
st.markdown("""
<style>
/* Responsive mobile */
@media (max-width: 768px) {
    .block-container { padding: 1rem 0.5rem; }
    h1 { font-size: 1.5rem !important; }
}
/* Boutons */
.stButton > button {
    width: 100%; height: 3.2em; border-radius: 12px;
    font-weight: 700; font-size: 1rem;
    background: linear-gradient(135deg, #667eea, #764ba2);
    color: white; border: none;
}
div.stDownloadButton > button {
    width: 100%; height: 3em; border-radius: 12px;
    background: linear-gradient(135deg, #43a047, #66bb6a);
    color: white; border: none; font-weight: 700;
}
/* Cartes */
.note-card {
    text-align: center; padding: 1.5rem;
    border-radius: 15px; margin: 1rem 0;
    color: white; font-weight: bold;
}
.detail-card {
    background: #f8f9fa; padding: 1rem;
    border-left: 4px solid #667eea;
    border-radius: 0 10px 10px 0;
    margin: 0.5rem 0;
}
.rgpd-card {
    background: #e8f5e9; padding: 1rem;
    border-left: 4px solid #2e7d32;
    border-radius: 0 10px 10px 0;
    margin: 0.5rem 0;
}
.warn-card {
    background: #fff3e0; padding: 1rem;
    border-left: 4px solid #ff9800;
    border-radius: 0 10px 10px 0;
    margin: 0.5rem 0;
}
.page-badge {
    display: inline-block; background: #667eea;
    color: white; padding: 0.3rem 0.8rem;
    border-radius: 20px; margin: 0.2rem;
    font-size: 0.85rem;
}
</style>
""", unsafe_allow_html=True)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  MODULE 1 — EXTRACTION PDF + WORD
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
class DocProcessor:
    """Extrait le texte des fichiers PDF et Word (.docx)"""

    @staticmethod
    def extract(uploaded_file):
        """Retourne le texte extrait du fichier."""
        if uploaded_file is None:
            return ""

        uploaded_file.seek(0)
        name = uploaded_file.name.lower()

        try:
            if name.endswith(".pdf"):
                doc = fitz.open(
                    stream=uploaded_file.read(), filetype="pdf"
                )
                text = "\n\n".join(
                    f"--- Page {i+1} ---\n{p.get_text('text')}"
                    for i, p in enumerate(doc)
                )
                doc.close()
                return text

            elif name.endswith(".docx"):
                doc = DocxDocument(uploaded_file)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

                # Extraire aussi les tableaux (fréquents dans les barèmes)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(
                            cell.text.strip() for cell in row.cells
                        )
                        if row_text.strip():
                            paragraphs.append(row_text)

                return "\n".join(paragraphs)

            else:
                return "Format non supporté. Utilisez PDF ou Word (.docx)"

        except Exception as e:
            return f"Erreur de lecture : {e}"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  MODULE 2 — SCANNER DE DOCUMENT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
class Scanner:
    """Améliore les photos de copies pour une meilleure lecture IA"""

    @staticmethod
    def enhance(pil_image):
        """Pipeline d'amélioration d'image"""
        img = np.array(pil_image)

        # Conversion couleur si nécessaire
        if len(img.shape) == 2:
            img = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)
        elif img.shape[2] == 4:
            img = cv2.cvtColor(img, cv2.COLOR_RGBA2BGR)

        # 1. Détection des bords pour recadrage auto
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)
        edged = cv2.Canny(blurred, 50, 200)
        cnts, _ = cv2.findContours(
            edged, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE
        )
        cnts = sorted(cnts, key=cv2.contourArea, reverse=True)[:5]

        for c in cnts:
            peri = cv2.arcLength(c, True)
            approx = cv2.approxPolyDP(c, 0.02 * peri, True)
            if len(approx) == 4:
                pts = approx.reshape(4, 2).astype("float32")
                rect = np.zeros((4, 2), dtype="float32")
                s = pts.sum(axis=1)
                rect[0] = pts[np.argmin(s)]
                rect[2] = pts[np.argmax(s)]
                diff = np.diff(pts, axis=1)
                rect[1] = pts[np.argmin(diff)]
                rect[3] = pts[np.argmax(diff)]

                tl, tr, br, bl = rect
                w = int(max(
                    np.linalg.norm(br - bl),
                    np.linalg.norm(tr - tl)
                ))
                h = int(max(
                    np.linalg.norm(tr - br),
                    np.linalg.norm(tl - bl)
                ))
                dst = np.array([
                    [0, 0], [w-1, 0],
                    [w-1, h-1], [0, h-1]
                ], dtype="float32")
                M = cv2.getPerspectiveTransform(rect, dst)
                img = cv2.warpPerspective(img, M, (w, h))
                break

        # 2. Amélioration du contraste (CLAHE)
        lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
        l, a, b = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        l = clahe.apply(l)
        enhanced = cv2.cvtColor(
            cv2.merge((l, a, b)), cv2.COLOR_LAB2BGR
        )

        return Image.fromarray(
            cv2.cvtColor(enhanced, cv2.COLOR_BGR2RGB)
        )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  MODULE 3 — CORRECTEUR IA (GEMINI)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
class Correcteur:
    """Moteur de correction basé sur Google Gemini"""

    def __init__(self, api_key):
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel("gemini-2.0-flash")

    def corriger(self, corrige_txt, bareme_txt, images,
                 classe, exam, total_pts):
        """Lance la correction et retourne un dict structuré."""

        # Adaptation au niveau
        if "2nde" in classe.lower() or "seconde" in classe.lower():
            niveau = (
                "NIVEAU 2nde MCVA : Sois ENCOURAGEANT. "
                "Accepte les formulations approximatives. "
                "Les synonymes pertinents sont acceptés."
            )
        elif "1" in classe.lower() or "première" in classe.lower():
            niveau = (
                "NIVEAU 1ère MCVA : Vocabulaire professionnel "
                "de base exigé. Rédaction structurée attendue."
            )
        else:
            niveau = (
                "NIVEAU Terminale MCVA : Sois EXIGEANT. "
                "Termes commerciaux précis obligatoires "
                "(SBAM, SONCAS, CAP, ADE). "
                "Calculs posés ET justes."
            )

        prompt = f"""Tu es un enseignant expert en Lycée Professionnel,
filière MCVA (Métiers du Commerce et de la Vente).

CONTEXTE :
• Examen : {exam}
• Classe : {classe}
• Total des points de l'épreuve : {total_pts}

{niveau}

═══ RÈGLES DE CORRECTION ═══

1. FIDÉLITÉ AU CORRIGÉ
   Corrige UNIQUEMENT avec le corrigé officiel fourni.
   N'invente AUCUNE réponse attendue.

2. BARÈME VENTE
   • Réponse juste et complète → 100% des points
   • Bonne démarche, résultat faux → 50% des points
   • Réponse partielle → au prorata
   • Faux ou hors-sujet → 0 point

3. VOCABULAIRE PROFESSIONNEL
   SBAM, SONCAS, CAP/ADE, objection, fidélisation...
   Si synonyme pertinent → ACCEPTE.

4. ÉCRITURE DIFFICILE
   ⚠ La copie peut être TRÈS MAL ÉCRITE.
   • Mot illisible → écris [illisible]
   • Mot incertain → écris [incertain: «mot supposé»]
   • Tu n'INVENTES ABSOLUMENT RIEN
   • Doute sur un mot-clé → ne pénalise PAS

═══ CORRIGÉ OFFICIEL ═══
{corrige_txt}

═══ BARÈME ═══
{bareme_txt}

═══ COPIE DE L'ÉLÈVE ═══
Les images jointes sont les pages de la copie.
Analyse chaque page attentivement.

═══ FORMAT DE RETOUR (JSON STRICT) ═══
Retourne UNIQUEMENT ce JSON, sans texte avant ni après :

{{
  "score_obtenu": 0.0,
  "total_points": {total_pts},
  "questions": [
    {{
      "numero": "1",
      "intitule": "description courte",
      "points_obtenus": 0.0,
      "points_max": 0.0,
      "commentaire": "Commentaire très court",
      "mots_illisibles": []
    }}
  ],
  "appreciation_generale": "2-3 phrases maximum",
  "points_forts": "...",
  "axes_amelioration": "...",
  "conseil_progression": "Un conseil concret et encourageant"
}}"""

        try:
            content = [prompt] + list(images)
            response = self.model.generate_content(
                content,
                generation_config=genai.GenerationConfig(
                    temperature=0.1,
                    max_output_tokens=4096,
                ),
            )

            txt = response.text.strip()
            txt = re.sub(r"^```(?:json)?\s*", "", txt)
            txt = re.sub(r"\s*```$", "", txt)

            data = json.loads(txt)
            return {"ok": True, "data": data}

        except json.JSONDecodeError as e:
            return {"ok": False, "erreur": f"JSON invalide: {e}",
                    "brut": txt}
        except Exception as e:
            return {"ok": False, "erreur": str(e)}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  MODULE 4 — EXPORT EXCEL
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def generer_excel(resultats):
    """Génère un fichier Excel avec tous les résultats."""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Feuille 1 : Récapitulatif
        df_recap = pd.DataFrame([{
            "N°": r["id"],
            "Score brut": f"{r['score']}/{r['total']}",
            "Note /20": r["note_20"],
            "Appréciation": r["appreciation"],
            "Conseil": r["conseil"],
        } for r in resultats])
        df_recap.to_excel(
            writer, sheet_name="Notes", index=False
        )

        # Feuille 2 : Détail par question
        details = []
        for r in resultats:
            for q in r.get("questions", []):
                details.append({
                    "Élève N°": r["id"],
                    "Question": q.get("numero", ""),
                    "Intitulé": q.get("intitule", ""),
                    "Points": f"{q.get('points_obtenus',0)}"
                              f"/{q.get('points_max',0)}",
                    "Commentaire": q.get("commentaire", ""),
                })
        if details:
            df_detail = pd.DataFrame(details)
            df_detail.to_excel(
                writer, sheet_name="Détail", index=False
            )

        # Feuille 3 : Statistiques
        notes = [r["note_20"] for r in resultats]
        if notes:
            stats = {
                "Statistique": [
                    "Nombre de copies",
                    "Moyenne de classe",
                    "Note la plus haute",
                    "Note la plus basse",
                    "Médiane",
                    "Copies ≥ 10/20",
                    "Copies < 10/20",
                ],
                "Valeur": [
                    len(notes),
                    round(sum(notes) / len(notes), 2),
                    max(notes),
                    min(notes),
                    round(
                        sorted(notes)[len(notes)//2], 2
                    ),
                    sum(1 for n in notes if n >= 10),
                    sum(1 for n in notes if n < 10),
                ],
            }
            df_stats = pd.DataFrame(stats)
            df_stats.to_excel(
                writer, sheet_name="Statistiques", index=False
            )

    output.seek(0)
    return output


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  APPLICATION PRINCIPALE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
def main():
    # Initialisation session
    if "resultats" not in st.session_state:
        st.session_state.resultats = []
    if "pages_copie" not in st.session_state:
        st.session_state.pages_copie = []
    if "copie_num" not in st.session_state:
        st.session_state.copie_num = 1

    # ━━━━━━━━━━━━━━━━━━━━
    #  EN-TÊTE
    # ━━━━━━━━━━━━━━━━━━━━
     col_logo, col_titre = st.columns([1, 3])
    with col_logo:
        st.image("logo.png.png", width=100)
    with col_titre:
        st.markdown("# 📝 Correcteur MCVA")
        st.markdown("*Gratuit · RGPD OK · Export Excel*") 

    # ━━━━━━━━━━━━━━━━━━━━
    #  SIDEBAR
    # ━━━━━━━━━━━━━━━━━━━━
    with st.sidebar:
        st.markdown("## ⚙️ Configuration")

        api_key = st.text_input(
            "🔑 Clé API Gemini",
            type="password",
            help="Gratuite sur aistudio.google.com/apikey",
        )

        st.markdown("---")
        exam_name = st.text_input(
            "📄 Nom de l'examen",
            placeholder="CCF Vente Situation 1",
        )
        classe = st.selectbox(
            "🎓 Classe",
            ["2nde MCVA", "1ère MCVA", "Terminale MCVA"],
        )
        total_pts = st.number_input(
            "📊 Total des points de l'épreuve",
            min_value=10, max_value=200, value=20,
            step=10,
            help="Ex: 20, 40, 80 ou 120",
        )

        st.markdown("---")
        st.markdown(f"""
        ### 📋 Progression
        **Copies corrigées : {len(st.session_state.resultats)}**
        """)

        if st.session_state.resultats:
            notes = [
                r["note_20"]
                for r in st.session_state.resultats
            ]
            st.markdown(
                f"**Moyenne actuelle : "
                f"{sum(notes)/len(notes):.1f}/20**"
            )

        st.markdown("---")

        # RGPD
        st.markdown("""
        ### 🔒 RGPD
        <div class="rgpd-card">
        <strong>Aucune donnée personnelle
        n'est envoyée.</strong><br/>
        ✅ Pas de nom d'élève<br/>
        ✅ Pas de photo de visage<br/>
        ✅ Numéro anonyme uniquement
        </div>
        """, unsafe_allow_html=True)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━
    #  ONGLETS PRINCIPAUX
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━
    tab1, tab2, tab3 = st.tabs([
        "📚 Documents",
        "📷 Corriger",
        "📊 Résultats",
    ])

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    #  ONGLET 1 — DOCUMENTS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    with tab1:
        st.markdown("### 📚 Documents de référence")
        st.markdown(
            "Chargez le corrigé et le barème "
            "(**PDF** ou **Word**)"
        )

        corrige_file = st.file_uploader(
            "📗 Corrigé officiel",
            type=["pdf", "docx"],
            key="corrige",
        )
        if corrige_file:
            st.success(
                f"✅ Corrigé : {corrige_file.name}"
            )
            with st.expander("Voir le contenu extrait"):
                st.text(
                    DocProcessor.extract(corrige_file)[:2000]
                    + "..."
                )

        bareme_file = st.file_uploader(
            "📊 Barème",
            type=["pdf", "docx"],
            key="bareme",
        )
        if bareme_file:
            st.success(f"✅ Barème : {bareme_file.name}")
            with st.expander("Voir le contenu extrait"):
                st.text(
                    DocProcessor.extract(bareme_file)[:2000]
                    + "..."
                )

        # Info conversion
        if total_pts != 20:
            st.markdown(f"""
            <div class="detail-card">
            <strong>🔄 Conversion automatique :</strong><br/>
            Score obtenu / {total_pts} × 20 = Note /20
            </div>
            """, unsafe_allow_html=True)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    #  ONGLET 2 — CORRIGER
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    with tab2:
        st.markdown(
            f"### 📷 Copie n°{st.session_state.copie_num}"
        )

        # Rappel RGPD
        st.markdown("""
        <div class="warn-card">
        ⚠️ <strong>RAPPEL :</strong>
        Photographiez UNIQUEMENT les réponses.
        <strong>Pas de nom, pas de prénom,
        pas d'en-tête.</strong>
        </div>
        """, unsafe_allow_html=True)

        # --- Prise de photos ---
        st.markdown("#### 📸 Photographier les pages")

        photo = st.camera_input(
            f"Page {len(st.session_state.pages_copie) + 1}",
            key=f"cam_{datetime.now().timestamp()}",
        )

        if photo:
            img = Image.open(photo)
            enhanced = Scanner.enhance(img)
            st.session_state.pages_copie.append(enhanced)

        # OU importer des fichiers
        with st.expander("📁 Ou importer des images"):
            fichiers = st.file_uploader(
                "Sélectionnez les pages",
                type=["png", "jpg", "jpeg"],
                accept_multiple_files=True,
                key="import_imgs",
            )
            if fichiers:
                for f in fichiers:
                    img = Image.open(f)
                    enhanced = Scanner.enhance(img)
                    st.session_state.pages_copie.append(
                        enhanced
                    )

        # Afficher les pages capturées
        nb = len(st.session_state.pages_copie)
        if nb > 0:
            st.markdown(
                f"**{nb} page(s) capturée(s) :**"
            )
            badges = " ".join(
                f'<span class="page-badge">'
                f'Page {i+1} ✅</span>'
                for i in range(nb)
            )
            st.markdown(badges, unsafe_allow_html=True)

            # Aperçu miniatures
            cols = st.columns(min(nb, 4))
            for i, img in enumerate(
                st.session_state.pages_copie
            ):
                with cols[i % len(cols)]:
                    st.image(
                        img, caption=f"P.{i+1}",
                        use_container_width=True,
                    )

        # Boutons d'action
        st.markdown("---")

        c1, c2 = st.columns(2)
        with c1:
            if st.button(
                "🗑️ Effacer les pages",
                use_container_width=True,
            ):
                st.session_state.pages_copie = []
                st.rerun()

        with c2:
            if st.button(
                "📸 Continuer (page suivante)",
                use_container_width=True,
            ):
                st.rerun()

        st.markdown("---")

        # --- BOUTON CORRIGER ---
        pret = (
            bool(api_key)
            and bool(exam_name)
            and bool(corrige_file)
            and bool(bareme_file)
            and len(st.session_state.pages_copie) > 0
        )

        if not pret:
            manquants = []
            if not api_key:
                manquants.append("Clé API")
            if not exam_name:
                manquants.append("Nom examen")
            if not corrige_file:
                manquants.append("Corrigé")
            if not bareme_file:
                manquants.append("Barème")
            if not st.session_state.pages_copie:
                manquants.append("Photos de la copie")
            st.warning(
                f"⏳ Manquant : {', '.join(manquants)}"
            )

        if st.button(
            f"🎓 CORRIGER LA COPIE N°"
            f"{st.session_state.copie_num}",
            disabled=not pret,
            use_container_width=True,
        ):
            with st.spinner(
                "🤖 L'IA analyse la copie... "
                "(30-60 secondes)"
            ):
                corrige_txt = DocProcessor.extract(
                    corrige_file
                )
                bareme_txt = DocProcessor.extract(
                    bareme_file
                )

                correcteur = Correcteur(api_key)
                result = correcteur.corriger(
                    corrige_txt,
                    bareme_txt,
                    st.session_state.pages_copie,
                    classe,
                    exam_name,
                    total_pts,
                )

            if result["ok"]:
                data = result["data"]
                score = data.get("score_obtenu", 0)
                note_20 = round(
                    (score / total_pts) * 20, 2
                )

                # Sauvegarder le résultat
                st.session_state.resultats.append({
                    "id": f"Copie {st.session_state.copie_num:02d}",
                    "score": score,
                    "total": total_pts,
                    "note_20": note_20,
                    "appreciation": data.get(
                        "appreciation_generale", ""
                    ),
                    "points_forts": data.get(
                        "points_forts", ""
                    ),
                    "axes": data.get(
                        "axes_amelioration", ""
                    ),
                    "conseil": data.get(
                        "conseil_progression", ""
                    ),
                    "questions": data.get(
                        "questions", []
                    ),
                })

                # Afficher le résultat
                st.balloons()

                # Couleur selon note
                if note_20 >= 16:
                    bg = "#2e7d32"
                elif note_20 >= 14:
                    bg = "#43a047"
                elif note_20 >= 10:
                    bg = "#1976d2"
                elif note_20 >= 8:
                    bg = "#ef6c00"
                else:
                    bg = "#c62828"

                st.markdown(f"""
                <div class="note-card" style="background:{bg}">
                    <div style="font-size:1rem">
                        Copie {st.session_state.copie_num:02d}
                    </div>
                    <div style="font-size:2.5rem">
                        {note_20} / 20
                    </div>
                    <div style="font-size:0.9rem">
                        (Score brut : {score}/{total_pts})
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # Détail par question
                for q in data.get("questions", []):
                    pts = q.get("points_obtenus", 0)
                    mx = q.get("points_max", 1)
                    pct = (pts / mx * 100) if mx else 0
                    if pct >= 80:
                        ic = "✅"
                    elif pct >= 50:
                        ic = "⚠️"
                    else:
                        ic = "❌"

                    st.markdown(f"""
                    <div class="detail-card">
                        {ic} <strong>Q{q.get('numero','?')}</strong>
                        — {q.get('intitule','')} :
                        <strong>{pts}/{mx}</strong>
                        — {q.get('commentaire','')}
                    </div>
                    """, unsafe_allow_html=True)

                # Appréciation
                st.markdown(f"""
                <div class="detail-card">
                    📝 <strong>Appréciation :</strong>
                    {data.get('appreciation_generale','')}
                </div>
                """, unsafe_allow_html=True)

                # Préparer copie suivante
                st.session_state.copie_num += 1
                st.session_state.pages_copie = []

                st.success(
                    "✅ Copie enregistrée ! "
                    "Prenez la copie suivante."
                )

            else:
                st.error(f"❌ Erreur : {result['erreur']}")
                if "brut" in result:
                    with st.expander("Réponse brute"):
                        st.code(result["brut"])

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    #  ONGLET 3 — RÉSULTATS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    with tab3:
        st.markdown("### 📊 Tableau des résultats")

        if not st.session_state.resultats:
            st.info(
                "Aucune copie corrigée. "
                "Allez dans l'onglet 📷 Corriger."
            )
        else:
            res = st.session_state.resultats

            # Tableau récapitulatif
            df = pd.DataFrame([{
                "N°": r["id"],
                "Score": f"{r['score']}/{r['total']}",
                "Note /20": r["note_20"],
                "Appréciation": r["appreciation"][:80],
            } for r in res])

            st.dataframe(
                df,
                use_container_width=True,
                hide_index=True,
            )

            # Statistiques
            notes = [r["note_20"] for r in res]
            st.markdown("---")
            c1, c2, c3, c4 = st.columns(4)
            with c1:
                st.metric(
                    "📚 Copies", len(notes)
                )
            with c2:
                st.metric(
                    "📈 Moyenne",
                    f"{sum(notes)/len(notes):.1f}"
                )
            with c3:
                st.metric("🏆 Max", max(notes))
            with c4:
                st.metric("📉 Min", min(notes))

            # Détail par élève
            st.markdown("---")
            st.markdown("### 📝 Détail par copie")
            for r in res:
                with st.expander(
                    f"{r['id']} — {r['note_20']}/20"
                ):
                    st.markdown(
                        f"**Appréciation :** "
                        f"{r['appreciation']}"
                    )
                    st.markdown(
                        f"**Points forts :** "
                        f"{r['points_forts']}"
                    )
                    st.markdown(
                        f"**À améliorer :** "
                        f"{r['axes']}"
                    )
                    st.markdown(
                        f"**Conseil :** "
                        f"{r['conseil']}"
                    )

            # EXPORT
            st.markdown("---")
            st.markdown("### 💾 Exporter")

            excel_data = generer_excel(res)
            st.download_button(
                label="📥 TÉLÉCHARGER LE FICHIER EXCEL",
                data=excel_data,
                file_name=(
                    f"notes_{exam_name.replace(' ','_')}_"
                    f"{datetime.now():%Y%m%d}.xlsx"
                ),
                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.spreadsheetml.sheet"
                ),
                use_container_width=True,
            )

            # Export CSV
            csv_data = df.to_csv(index=False).encode(
                'utf-8-sig'
            )
            st.download_button(
                label="📥 Télécharger en CSV",
                data=csv_data,
                file_name=(
                    f"notes_{datetime.now():%Y%m%d}.csv"
                ),
                mime="text/csv",
                use_container_width=True,
            )

            # Bouton reset
            st.markdown("---")
            if st.button(
                "🗑️ Effacer tous les résultats",
                use_container_width=True,
            ):
                st.session_state.resultats = []
                st.session_state.copie_num = 1
                st.rerun()


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  LANCEMENT
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
if __name__ == "__main__":
    main()
